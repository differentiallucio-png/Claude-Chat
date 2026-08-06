#!/usr/bin/env python3
"""SOPHIA DOCX Analyzer 0.3.0.

Analisador descritivo e auditável para manuscritos acadêmicos em DOCX.
Não atribui autoria por IA, não produz nota global de qualidade e não usa
faixas SEC universais. Toda métrica inclui proveniência, escopo e limitações.

Uso:
  python sophia_docx_analyzer.py manuscrito.docx --out-dir saida
"""
from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import statistics
import sys
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Cm

VERSION = "0.3.0"
TOKENIZER_NAME = "sophia-word-regex-pt-v1"
SENTENCE_SEGMENTER = "sophia-paragraph-regex-v1"

WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9]+(?:[-'’][A-Za-zÀ-ÖØ-öø-ÿ0-9]+)*", re.UNICODE)
YEAR_RE = re.compile(r"\b(?:18|19|20)\d{2}[a-z]?\b")
PAREN_CIT_RE = re.compile(r"\([^()]{0,180}\b(?:18|19|20)\d{2}[a-z]?[^()]{0,80}\)")
NARRATIVE_CIT_RE = re.compile(r"\b[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ][A-Za-zÁÀÂÃÉÊÍÓÔÕÚÇáàâãéêíóôõúç'’\-]+(?:\s+et\s+al\.)?\s*\((?:18|19|20)\d{2}[a-z]?\)")
PAGE_CIT_RE = re.compile(r"\b(?:p\.|pp\.)\s*\d+(?:\s*[-–]\s*\d+)?", re.I)
INLINE_QUOTE_RE = re.compile(r"[“\"]([^“”\"]{3,700})[”\"]")
SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+(?=(?:[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ0-9“\"]|\([A-ZÁÀÂÃÉÊÍÓÔÕÚÇ]))")

HEADING_STYLE_RE = re.compile(r"heading|t[ií]tulo|sectionheading|cabe[cç]alho", re.I)
REFERENCE_STYLE_RE = re.compile(r"refer[eê]ncia", re.I)
BLOCK_QUOTE_STYLE_RE = re.compile(r"cita[cç][aã]o|quote|blockquote", re.I)
DECLARATION_STYLE_RE = re.compile(r"declara[cç][aã]o", re.I)

PARAGRAPH_CONNECTORS = [
    "além disso", "ademais", "contudo", "no entanto", "todavia", "portanto",
    "por conseguinte", "nesse sentido", "dessa forma", "desse modo", "assim",
    "com efeito", "de resto", "em contrapartida", "por outro lado", "ainda assim",
    "entretanto", "logo", "por isso", "por essa razão", "em consequência",
    "ao contrário", "em contraste", "apesar disso", "não obstante", "em primeiro lugar",
    "em segundo lugar", "primeiro", "segundo", "terceiro", "quarto", "por fim",
    "antes de", "depois de", "para compreender", "para evitar", "para fins de",
    "nessa direção", "sob esse aspecto", "diante disso", "a partir disso",
]

META_PATTERNS = [
    r"\beste ensaio\b", r"\bo presente ensaio\b", r"\bneste texto\b",
    r"\ba contribuição pretendida\b", r"\bpara fins deste ensaio\b",
    r"\bcomo veremos\b", r"\bexaminado adiante\b", r"\bdiscutido adiante\b",
    r"\bo que foi articulado antes\b", r"\besta seção\b", r"\bna seção anterior\b",
    r"\bpropõe-se\b", r"\bsustenta-se\b", r"\bdenomina-se\b",
]
META_RE = [re.compile(p, re.I) for p in META_PATTERNS]

RELATION_MARKERS = [
    "converge", "convergência", "diverge", "divergência", "difere", "diferença",
    "complementa", "complementares", "em cadeia", "não dizem a mesma coisa",
    "em planos diferentes", "em escalas diferentes", "ao passo que", "enquanto",
    "em contraste", "não competem", "se articulam sem se confundirem", "por sua vez",
    "primeiro", "segundo", "terceiro", "o primeiro", "o segundo",
]

NOMINAL_SUFFIXES = {
    "acao_sao": re.compile(r"(?:ção|ções|são|sões)$", re.I),
    "mento": re.compile(r"mento(?:s)?$", re.I),
    "dade": re.compile(r"dade(?:s)?$", re.I),
    "encia_ancia": re.compile(r"(?:ência|ências|ância|âncias)$", re.I),
}

@dataclass
class ParagraphRecord:
    index: int
    style: str
    text: str
    section: str
    is_heading: bool
    is_reference: bool
    is_declaration: bool
    is_block_quote: bool
    in_main_body: bool
    in_abstract: bool


def words(text: str) -> list[str]:
    return WORD_RE.findall(text or "")


def word_count(text: str) -> int:
    return len(words(text))


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def split_sentences(text: str) -> list[str]:
    text = normalize(text)
    if not text:
        return []
    parts = [p.strip() for p in SENTENCE_SPLIT_RE.split(text) if p.strip()]
    return parts or [text]


def style_name(paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def is_heading_style(style: str) -> bool:
    return bool(HEADING_STYLE_RE.search(style or ""))


def is_reference_style(style: str) -> bool:
    return bool(REFERENCE_STYLE_RE.search(style or ""))


def is_declaration_style(style: str) -> bool:
    return bool(DECLARATION_STYLE_RE.search(style or ""))


def has_quote_indentation(paragraph) -> bool:
    fmt = paragraph.paragraph_format
    left = fmt.left_indent
    right = fmt.right_indent
    # 0.75 cm is conservative: catches block quotations but avoids ordinary first-line indent.
    return bool((left and left >= Cm(0.75)) or (right and right >= Cm(0.75)))


def is_block_quote(paragraph) -> bool:
    style = style_name(paragraph)
    text = normalize(paragraph.text)
    if BLOCK_QUOTE_STYLE_RE.search(style):
        return True
    if len(words(text)) >= 30 and has_quote_indentation(paragraph) and not is_heading_style(style):
        return True
    return False


def table_texts(doc: Document) -> list[dict[str, Any]]:
    out = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = [normalize(cell.text) for cell in row.cells]
            rows.append({"row": ri, "cells": cells, "word_count": sum(word_count(c) for c in cells)})
        out.append({"table": ti, "rows": rows, "word_count": sum(r["word_count"] for r in rows)})
    return out


HEADING_NUMBER_RE = re.compile(r"^\s*\d+(\.\d+)*\s*[.)\-]?\s*")


def strip_heading_number(text: str) -> str:
    """Remove numeração de título (ex.: '1 Introdução' -> 'Introdução', '2.3 Método' -> 'Método').

    Sem isso, headings.casefold() nunca bate 'introdução' em manuscritos que numeram
    títulos de nível 1 (padrão comum em ABNT), e todo o corpo é lido como zero (SYS-01).
    """
    return HEADING_NUMBER_RE.sub("", text or "").strip()


def build_records(doc: Document) -> list[ParagraphRecord]:
    headings = [(i, normalize(p.text)) for i, p in enumerate(doc.paragraphs) if is_heading_style(style_name(p))]
    intro_idx = next((i for i, t in headings if strip_heading_number(t).casefold() == "introdução"), None)
    refs_idx = next((i for i, t in headings if strip_heading_number(t).casefold() == "referências"), None)
    abstract_idx = next((i for i, t in headings if strip_heading_number(t).casefold() == "resumo"), None)
    current_section = "PRÉ-TEXTO"
    records = []
    for i, p in enumerate(doc.paragraphs):
        text = normalize(p.text)
        style = style_name(p)
        heading = is_heading_style(style)
        if heading and text:
            current_section = text
        main = bool(intro_idx is not None and refs_idx is not None and intro_idx <= i < refs_idx)
        abstract = bool(abstract_idx is not None and intro_idx is not None and abstract_idx < i < intro_idx)
        records.append(ParagraphRecord(
            index=i,
            style=style,
            text=text,
            section=current_section,
            is_heading=heading,
            is_reference=is_reference_style(style) or (refs_idx is not None and i > refs_idx and not is_declaration_style(style)),
            is_declaration=is_declaration_style(style),
            is_block_quote=is_block_quote(p),
            in_main_body=main,
            in_abstract=abstract,
        ))
    return records


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def docx_xml_features(path: Path) -> dict[str, Any]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    result = {"citation_fields": 0, "field_codes": [], "styles_seen": []}
    with zipfile.ZipFile(path) as zf:
        if "word/document.xml" in zf.namelist():
            root = ET.fromstring(zf.read("word/document.xml"))
            codes = []
            for node in root.findall(".//w:instrText", ns):
                if node.text:
                    codes.append(node.text.strip())
            result["field_codes"] = codes
            result["citation_fields"] = sum(1 for c in codes if "CITATION" in c.upper())
        if "word/styles.xml" in zf.namelist():
            sroot = ET.fromstring(zf.read("word/styles.xml"))
            names = []
            for st in sroot.findall(".//w:style", ns):
                sid = st.attrib.get(f"{{{ns['w']}}}styleId", "")
                n = st.find("w:name", ns)
                val = n.attrib.get(f"{{{ns['w']}}}val", "") if n is not None else ""
                names.append({"style_id": sid, "name": val})
            result["styles_seen"] = names
    return result


def provenance(name: str, scope: str, unit: str, denominator: str, method: str, limitations: list[str], coverage: dict[str, Any] | None = None) -> dict[str, Any]:
    return {
        "metric": name,
        "scope": scope,
        "unit": unit,
        "tokenizer": TOKENIZER_NAME,
        "sentence_segmenter": SENTENCE_SEGMENTER,
        "analyzer_version": VERSION,
        "denominator": denominator,
        "method": method,
        "coverage": coverage or {},
        "known_limitations": limitations,
    }


def nominal_candidates(text: str) -> dict[str, Any]:
    excluded = {"são"}
    toks = [t.casefold() for t in words(text) if not t.isdigit() and t.casefold() not in excluded]
    by_kind: dict[str, Counter[str]] = {k: Counter() for k in NOMINAL_SUFFIXES}
    for tok in toks:
        for kind, rx in NOMINAL_SUFFIXES.items():
            if rx.search(tok):
                by_kind[kind][tok] += 1
                break
    total = sum(sum(c.values()) for c in by_kind.values())
    return {
        "label": "candidatos_por_sufixo_nao_confirmados",
        "total_candidates": total,
        "per_1000_words": (total / len(toks) * 1000) if toks else 0.0,
        "by_suffix": {k: dict(v.most_common(30)) for k, v in by_kind.items()},
        "warning": "Sem análise morfossintática contextual, estes itens não devem ser chamados de nominalizações confirmadas nem classificados como técnicos ou ornamentais.",
    }


def paragraph_connector(text: str) -> str:
    low = normalize(text).casefold()
    for c in sorted(PARAGRAPH_CONNECTORS, key=len, reverse=True):
        if low == c or low.startswith(c + " ") or low.startswith(c + ","):
            return c
    return ""


def metadiscourse_hits(text: str) -> list[str]:
    hits = []
    for rx in META_RE:
        for m in rx.finditer(text):
            hits.append(m.group(0))
    return hits


def citation_paragraph_info(records: list[ParagraphRecord]) -> list[dict[str, Any]]:
    out = []
    for r in records:
        if not r.in_main_body or r.is_heading or not r.text:
            continue
        p_cits = PAREN_CIT_RE.findall(r.text)
        n_cits = NARRATIVE_CIT_RE.findall(r.text)
        if not (p_cits or n_cits or YEAR_RE.search(r.text)):
            continue
        markers = [m for m in RELATION_MARKERS if m in r.text.casefold()]
        out.append({
            "paragraph_index": r.index,
            "section": r.section,
            "style": r.style,
            "word_count": word_count(r.text),
            "parenthetical_citations": p_cits,
            "narrative_citations": n_cits,
            "relation_markers": markers,
            "preview": r.text[:500],
            "classification": "candidate_only",
            "note": "A presença de citação ou marcador não determina resumo, evidência ou comentário; requer leitura humana da função retórica.",
        })
    return out


def direct_quotes(records: list[ParagraphRecord]) -> dict[str, Any]:
    blocks = []
    inline = []
    for r in records:
        if not r.text:
            continue
        if r.is_block_quote:
            blocks.append({
                "paragraph_index": r.index,
                "section": r.section,
                "style": r.style,
                "word_count": word_count(r.text),
                "has_author_year": bool(YEAR_RE.search(r.text)),
                "has_page": bool(PAGE_CIT_RE.search(r.text)),
                "preview": r.text[:500],
            })
        for m in INLINE_QUOTE_RE.finditer(r.text):
            q = normalize(m.group(1))
            if word_count(q) >= 2:
                inline.append({
                    "paragraph_index": r.index,
                    "section": r.section,
                    "word_count": word_count(q),
                    "quote": q[:500],
                    "context_has_year": bool(YEAR_RE.search(r.text)),
                    "context_has_page": bool(PAGE_CIT_RE.search(r.text)),
                })
    return {"block_quotes": blocks, "inline_quotes": inline}


def scope_text(records: list[ParagraphRecord], predicate, include_headings: bool = False, include_quotes: bool = True) -> str:
    parts = []
    for r in records:
        if not predicate(r) or not r.text:
            continue
        if r.is_heading and not include_headings:
            continue
        if r.is_block_quote and not include_quotes:
            continue
        parts.append(r.text)
    return "\n".join(parts)


def analyze(path: Path) -> dict[str, Any]:
    doc = Document(path)
    records = build_records(doc)
    tables = table_texts(doc)
    xml_features = docx_xml_features(path)

    all_para = scope_text(records, lambda r: True, include_headings=True, include_quotes=True)
    main_prose = scope_text(records, lambda r: r.in_main_body, include_headings=False, include_quotes=True)
    main_prose_no_quotes = scope_text(records, lambda r: r.in_main_body, include_headings=False, include_quotes=False)
    main_with_headings = scope_text(records, lambda r: r.in_main_body, include_headings=True, include_quotes=True)
    references_text = scope_text(records, lambda r: r.is_reference and not r.is_declaration, include_headings=False, include_quotes=True)
    declarations_text = scope_text(records, lambda r: r.is_declaration, include_headings=False, include_quotes=True)
    abstract_text = scope_text(records, lambda r: r.in_abstract, include_headings=False, include_quotes=True)
    table_words = sum(t["word_count"] for t in tables)

    main_sentences = []
    paragraph_sentence_counts = []
    for r in records:
        if r.in_main_body and not r.is_heading and r.text:
            s = split_sentences(r.text)
            main_sentences.extend(s)
            paragraph_sentence_counts.append({"paragraph_index": r.index, "sentences": len(s), "words": word_count(r.text), "section": r.section})
    sentence_lengths = [word_count(s) for s in main_sentences if word_count(s)]

    connector_records = []
    main_paragraphs = [r for r in records if r.in_main_body and not r.is_heading and r.text]
    for r in main_paragraphs:
        c = paragraph_connector(r.text)
        if c:
            connector_records.append({"paragraph_index": r.index, "section": r.section, "connector": c, "preview": r.text[:180]})

    meta_records = []
    for r in records:
        if r.in_main_body and not r.is_heading and r.text:
            hits = metadiscourse_hits(r.text)
            if hits:
                meta_records.append({"paragraph_index": r.index, "section": r.section, "hits": hits, "preview": r.text[:260]})

    quotes = direct_quotes(records)
    citations = citation_paragraph_info(records)
    nom = nominal_candidates(main_prose_no_quotes)

    section_stats: dict[str, dict[str, Any]] = defaultdict(lambda: {"paragraphs": 0, "words": 0, "citation_paragraphs": 0, "block_quotes": 0})
    citation_idx = {c["paragraph_index"] for c in citations}
    quote_idx = {q["paragraph_index"] for q in quotes["block_quotes"]}
    for r in main_paragraphs:
        s = section_stats[r.section]
        s["paragraphs"] += 1
        s["words"] += word_count(r.text)
        if r.index in citation_idx:
            s["citation_paragraphs"] += 1
        if r.index in quote_idx:
            s["block_quotes"] += 1

    scopes = {
        "document_paragraphs_including_headings_excluding_tables": word_count(all_para),
        "document_including_tables": word_count(all_para) + table_words,
        "abstract_prose": word_count(abstract_text),
        "main_body_prose_including_block_quotes_excluding_headings_tables_references": word_count(main_prose),
        "main_body_prose_excluding_block_quotes_headings_tables_references": word_count(main_prose_no_quotes),
        "main_body_with_headings_excluding_tables_references": word_count(main_with_headings),
        "table_words": table_words,
        "main_body_prose_plus_tables": word_count(main_prose) + table_words,
        "references": word_count(references_text),
        "declarations": word_count(declarations_text),
    }

    sentence_count = len(sentence_lengths)
    main_words = scopes["main_body_prose_including_block_quotes_excluding_headings_tables_references"]
    mean_sentence = (sum(sentence_lengths) / sentence_count) if sentence_count else 0.0

    consistency = []
    consistency.append({
        "test": "mean_sentence_arithmetic",
        "passed": math.isclose(mean_sentence, main_words / sentence_count if sentence_count else 0.0, rel_tol=1e-9, abs_tol=1e-9),
        "observed": mean_sentence,
        "recomputed": main_words / sentence_count if sentence_count else 0.0,
    })
    consistency.append({
        "test": "scope_document_ge_main",
        "passed": scopes["document_including_tables"] >= scopes["main_body_prose_plus_tables"],
        "observed": scopes["document_including_tables"],
        "main": scopes["main_body_prose_plus_tables"],
    })
    consistency.append({
        "test": "quote_detector_nonzero_guard",
        "passed": len(quotes["block_quotes"]) > 0 or not any(BLOCK_QUOTE_STYLE_RE.search(s.get("style_id", "") + " " + s.get("name", "")) for s in xml_features.get("styles_seen", [])),
        "observed_block_quotes": len(quotes["block_quotes"]),
    })
    advisories = []
    advisories.append({
        "test": "connector_zero_guard",
        "flagged": not (bool(connector_records) or len(main_paragraphs) < 20),
        "observed_connector_openings": len(connector_records),
        "main_paragraphs": len(main_paragraphs),
        "interpretation": (
            "Zero conectores do léxico não é falha de detecção por si só: pode ser estilo autoral "
            "legítimo (abertura de parágrafo por sintagma nominal ou adjunto, sem conector fixo). "
            "Este teste é consultivo e NÃO rebaixa o estatuto canônico de outras métricas (contagem "
            "de palavras, sentenças, parágrafos) quando falha; exige apenas auditoria de léxico "
            "específica para conectores, isoladamente."
        ),
    })
    failed = [t for t in consistency if not t["passed"]]

    report = {
        "tool": {"name": "SOPHIA DOCX Analyzer", "version": VERSION, "generated_at_utc": datetime.now(timezone.utc).isoformat()},
        "input": {"path": str(path), "filename": path.name, "sha256": sha256(path), "bytes": path.stat().st_size},
        "document_structure": {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "styles": dict(Counter(r.style for r in records)),
            "sections": [r.text for r in records if r.is_heading and r.text],
            "xml": {"citation_fields": xml_features["citation_fields"]},
        },
        "word_counts": {
            "values": scopes,
            "provenance": provenance(
                "word_counts", "multiple explicit scopes", "regex tokens", "scope-specific",
                "Contagem por expressão regular sobre parágrafos DOCX; tabelas contadas separadamente.",
                ["Hifenização editorial e símbolos matemáticos podem ser tokenizados de modo diferente de processadores de texto.", "Contagens só são comparáveis quando escopo e tokenizador coincidem."],
                {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
            ),
        },
        "sentences": {
            "count": sentence_count,
            "mean_words": mean_sentence,
            "median_words": statistics.median(sentence_lengths) if sentence_lengths else 0.0,
            "stdev_words": statistics.pstdev(sentence_lengths) if len(sentence_lengths) > 1 else 0.0,
            "over_35": sum(1 for n in sentence_lengths if n > 35),
            "over_45": sum(1 for n in sentence_lengths if n > 45),
            "under_9": sum(1 for n in sentence_lengths if n < 9),
            "max_words": max(sentence_lengths) if sentence_lengths else 0,
            "provenance": provenance(
                "sentence_length", "main body prose including block quotes, excluding headings/tables/references", "sentences and regex tokens", "detected sentences",
                "Segmentação heurística por pontuação terminal dentro de cada parágrafo.",
                ["Abreviaturas, iniciais, listas numeradas e citações podem afetar a segmentação.", "A métrica é descritiva; comprimento não equivale a qualidade."],
            ),
        },
        "direct_quotes": {
            **quotes,
            "word_styles_detected": [s for s in dict(Counter(r.style for r in records)) if BLOCK_QUOTE_STYLE_RE.search(s)],
            "citation_fields_in_xml": xml_features["citation_fields"],
            "provenance": provenance(
                "direct_quote_detection", "whole document", "quote occurrences", "paragraphs and inline spans",
                "Detecção multicamada por estilo, recuo, aspas e campos CITATION do XML.",
                ["Nem todo texto entre aspas é citação de fonte; pode ser termo, pergunta ou fala.", "A atribuição e a conformidade bibliográfica exigem leitura humana."],
            ),
        },
        "paragraph_opening_connectors": {
            "count": len(connector_records),
            "paragraphs": len(main_paragraphs),
            "ratio": len(connector_records) / len(main_paragraphs) if main_paragraphs else 0.0,
            "examples": connector_records[:50],
            "zero_result_requires_lexicon_audit": not connector_records and len(main_paragraphs) >= 20,
            "provenance": provenance(
                "paragraph_opening_connectors", "main body prose", "paragraph openings", "main body prose paragraphs",
                "Correspondência com léxico explícito e extensível de conectores/organizadores na abertura.",
                ["Léxico não exaustivo; zero não prova ausência do fenômeno.", "Abertura sem conector pode ser coesa por progressão temática ou cadeia referencial."],
                {"lexicon_entries": len(PARAGRAPH_CONNECTORS)},
            ),
        },
        "metadiscourse": {
            "occurrences": sum(len(x["hits"]) for x in meta_records),
            "records": meta_records,
            "provenance": provenance(
                "metadiscourse_candidates", "main body prose", "pattern matches", "main body prose words",
                "Padrões explícitos de autorreferência estrutural e orientação do percurso.",
                ["Lista não exaustiva.", "Metadiscurso pode ser funcional em ensaio analítico-filosófico; frequência não é defeito automático."],
                {"patterns": META_PATTERNS},
            ),
        },
        "nominalization_candidates": {
            **nom,
            "provenance": provenance(
                "nominalization_candidates", "main body prose excluding block quotes", "suffix candidates", "regex word tokens",
                "Detecção por sufixos nominais frequentes em português.",
                ["Não há etiquetagem morfossintática; resultados são candidatos, não nominalizações confirmadas.", "O script não classifica automaticamente termos como técnicos, necessários ou ornamentais."],
            ),
        },
        "source_integration": {
            "policy": "No universal SEC percentages. Automated output is a candidate map for human rhetorical coding.",
            "citation_bearing_paragraphs": len(citations),
            "paragraph_records": citations,
            "section_summary": dict(section_stats),
            "provenance": provenance(
                "source_integration_candidate_map", "main body prose", "citation-bearing paragraphs and relation-marker candidates", "main body prose paragraphs",
                "Extração por padrões autor-ano, anos e marcadores explícitos de relação entre vozes.",
                ["Não classifica automaticamente resumo, evidência e comentário.", "Uma unidade pode realizar várias operações simultaneamente.", "Qualquer julgamento exige trecho, seção, função e confiança."],
            ),
        },
        "readiness": {
            "argumentative": "requires_human_assessment",
            "bibliographic": "not_audited",
            "normative": "not_audited",
            "documentary": "partially_described_not_audited",
            "journal_specific": "indeterminate_without_target_journal",
        },
        "consistency_tests": {"passed": not failed, "tests": consistency, "failed_count": len(failed)},
        "advisories": {"items": advisories, "flagged_count": len([a for a in advisories if a["flagged"]])},
        "guardrails": [
            "Não interpretar zero sem auditoria de cobertura do detector.",
            "Zero de aberturas de conector (connector_zero_guard) é consultivo: não rebaixa a canonicidade de contagens de palavras, sentenças ou parágrafos, que são medidas independentes.",
            "Não converter candidatos por sufixo em nominalizações confirmadas.",
            "Não usar faixas SEC universais para ensaio teórico.",
            "Não recomendar expansão sem verificar se a relação já foi explicitada em outro trecho.",
            "Não declarar prontidão para submissão a partir deste relatório isolado.",
        ],
    }
    return report


def render_markdown(data: dict[str, Any]) -> str:
    wc = data["word_counts"]["values"]
    sent = data["sentences"]
    q = data["direct_quotes"]
    conn = data["paragraph_opening_connectors"]
    meta = data["metadiscourse"]
    nom = data["nominalization_candidates"]
    src = data["source_integration"]
    lines = [
        "# SOPHIA — Relatório descritivo auditável",
        "",
        f"**Arquivo:** `{data['input']['filename']}`  ",
        f"**SHA-256:** `{data['input']['sha256']}`  ",
        f"**Analisador:** {data['tool']['name']} {data['tool']['version']}",
        "",
        "> Este relatório não produz nota global, não detecta autoria por IA e não autoriza reescrita automática.",
        "",
        "## 1. Escopos de contagem",
        "",
        "| Escopo | Palavras |",
        "|---|---:|",
    ]
    for k, v in wc.items():
        lines.append(f"| `{k}` | {v} |")
    lines += [
        "",
        "## 2. Sentenças",
        "",
        f"- Sentenças detectadas: **{sent['count']}**",
        f"- Média: **{sent['mean_words']:.2f}** palavras",
        f"- Mediana: **{sent['median_words']:.2f}**",
        f"- Acima de 35 palavras: **{sent['over_35']}**",
        f"- Acima de 45 palavras: **{sent['over_45']}**",
        "",
        "## 3. Citações diretas",
        "",
        f"- Blocos detectados: **{len(q['block_quotes'])}**",
        f"- Citações/trechos inline entre aspas: **{len(q['inline_quotes'])}**",
        f"- Campos CITATION no XML: **{q['citation_fields_in_xml']}**",
        "",
    ]
    for item in q["block_quotes"]:
        lines.append(f"- § `{item['section']}` · parágrafo {item['paragraph_index']} · estilo `{item['style']}` · {item['word_count']} palavras")
    lines += [
        "",
        "## 4. Aberturas e metadiscurso",
        "",
        f"- Aberturas com conectores reconhecidos: **{conn['count']}/{conn['paragraphs']}**",
        f"- O detector exige auditoria por resultado zero: **{conn['zero_result_requires_lexicon_audit']}**",
        f"- Candidatos a metadiscurso: **{meta['occurrences']}**",
        "",
        "## 5. Candidatos a nominalização",
        "",
        f"- Candidatos por sufixo: **{nom['total_candidates']}** ({nom['per_1000_words']:.2f}/mil)",
        f"- Estatuto: **{nom['label']}**",
        "- O script não os classifica como técnicos ou ornamentais.",
        "",
        "## 6. Integração de fontes",
        "",
        f"- Parágrafos com candidatos a citações: **{src['citation_bearing_paragraphs']}**",
        "- Não foram calculadas porcentagens SEC normativas.",
        "- A saída registra localização, trecho e marcadores para codificação humana.",
        "",
        "## 7. Testes internos",
        "",
        f"**Resultado geral:** {'APROVADO' if data['consistency_tests']['passed'] else 'FALHOU'}",
        "",
    ]
    for t in data["consistency_tests"]["tests"]:
        lines.append(f"- {'OK' if t['passed'] else 'FALHA'} — `{t['test']}`")
    lines += [
        "",
        "## 8. Prontidão",
        "",
    ]
    for k, v in data["readiness"].items():
        lines.append(f"- **{k}:** `{v}`")
    lines += ["", "## 9. Salvaguardas", ""]
    for g in data["guardrails"]:
        lines.append(f"- {g}")
    return "\n".join(lines) + "\n"


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    ap.add_argument("--out-dir", type=Path, required=True)
    args = ap.parse_args()
    if not args.docx.exists():
        print(f"Arquivo não encontrado: {args.docx}", file=sys.stderr)
        return 2
    args.out_dir.mkdir(parents=True, exist_ok=True)
    data = analyze(args.docx)
    (args.out_dir / "metrics.json").write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (args.out_dir / "report.md").write_text(render_markdown(data), encoding="utf-8")
    if not data["consistency_tests"]["passed"]:
        print("Falha em testes internos; relatório gerado, mas não deve ser interpretado.", file=sys.stderr)
        return 3
    print(json.dumps({
        "metrics": str(args.out_dir / "metrics.json"),
        "report": str(args.out_dir / "report.md"),
        "consistency": "passed",
    }, ensure_ascii=False))
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
